

from base import BaseScraper
import logging
import os
import re
import time
import random
import requests
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Optional
from urllib.parse import urlparse, urljoin

# =============================================================================
# DOCUMENT PARSING (Optionnel mais recommandé)
# =============================================================================
try:
    import PyPDF2
    from docx import Document as WordDocument
    import openpyxl
    import pandas as pd
    DOCUMENT_PARSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PARSING_AVAILABLE = False
    logging.warning("⚠️  Libraries de parsing non installées: PyPDF2, python-docx, openpyxl, pandas")

logger = logging.getLogger(__name__)


class WebSearchSerperScraper(BaseScraper):
    """
    Scraper utilisant Serper.dev API pour la recherche Google + extraction de contenu
    
    Fonctionnalités :
    ✅ Recherche Google via Serper API (résultats de qualité)
    ✅ Visite des pages pour extraire le contenu réel
    ✅ Téléchargement et parsing des documents (PDF, Word, Excel, CSV)
    ✅ Filtrage intelligent pour appels d'offre Côte d'Ivoire
    ✅ Stockage MongoDB avec métadonnées enrichies
    
    ⚠️  Nécessite une clé API Serper.dev (100 requêtes gratuites/mois)
    🔗 https://serper.dev/
    """
    
    def __init__(self):
        super().__init__('WebSearch Serper', 'https://google.com')
        
        # Clé API Serper (à configurer dans .env)
        self.api_key = os.getenv('SERPER_API_KEY', '')
        self.api_url = 'https://google.serper.dev/search'
        
        # Mots-clés optimisés pour la Côte d'Ivoire
        self.keywords = [
            "appel d'offre Côte d'Ivoire",
            "marché public Abidjan",
            "consultation entreprise CI",
            "avis d'appel d'offres Côte d'Ivoire",
            "tender Ivory Coast",
            "Reseau Informatique appel d'offre",
            "Consultation batiment Côte d'Ivoire",
            "Cyber sécurité marché public",
            "Microsoft appel d'offre CI",
            "Developpement d'application marché public",
            "Système Information appel d'offre",
            "Vente de matériel informatique Côte d'Ivoire",
            "Informatique consultation entreprise",
            "Acquisition équipement CI",
            "DAO Côte d'Ivoire",
            "DCE marché public",
        ]
        
        # Configuration
        self.max_results_per_keyword = int(os.getenv('SERPER_MAX_RESULTS', '15'))
        self.max_pages_to_visit = int(os.getenv('SERPER_MAX_PAGES', '50'))
        self.download_documents = os.getenv('SERPER_DOWNLOAD_DOCS', 'true').lower() == 'true'
        self.download_folder = Path(os.getenv('SERPER_DOWNLOAD_FOLDER', 'downloads/serper'))
        self.download_folder.mkdir(parents=True, exist_ok=True)
        
        # Délais pour éviter les blocages
        self.delay_between_requests = random.uniform(2, 5)
        self.delay_between_keywords = random.uniform(3, 8)
        
        # Filtres de contenu
        self.relevant_keywords = [
            'appel d\'offre', 'appel d\'offres', 'marché public', 'consultation',
            'avis d\'appel', 'tender', 'bid', 'rfp', 'soumission',
            'prestataire', 'fournisseur', 'contrat public', 'DAO', 'DCE',
            'dossier de consultation', 'cahier des charges'
        ]
        
        self.exclude_keywords = [
            'emploi', 'recrutement', 'stage', 'cdd', 'cdi', 'hiring',
            'offre d\'emploi', 'job', 'career', 'postulez', 'rejoignez'
        ]
        
        # Extensions de documents à télécharger
        self.document_extensions = [
            '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv',
            '.ppt', '.pptx', '.txt', '.rtf', '.odt', '.ods'
        ]
    
    def scrape(self) -> List[Dict]:
        """Scrape le web via Serper API pour trouver des appels d'offre"""
        logger.info(f"🔄 Démarrage WebSearch Serper avec {len(self.keywords)} mots-clés...")
        
        if not self.api_key:
            logger.error("❌ SERPER_API_KEY non configurée dans .env")
            return []
        
        all_offres = []
        visited_urls = set()
        pages_visited = 0
        api_calls = 0
        
        for i, keyword in enumerate(self.keywords, 1):
            if pages_visited >= self.max_pages_to_visit:
                logger.info(f"✅ Limite de pages atteinte ({self.max_pages_to_visit})")
                break
            
            if api_calls >= 100:  # Limite gratuite Serper
                logger.warning(f"⚠️  Limite API Serper atteinte (100 req/mois)")
                break
                
            logger.info(f"[{i}/{len(self.keywords)}] Recherche Serper: '{keyword}'")
            
            try:
                # Effectuer la recherche via Serper API
                search_results = self._search_serper(keyword, max_results=self.max_results_per_keyword)
                api_calls += 1
                
                if not search_results:
                    logger.warning(f"   ⚠️  Aucun résultat pour '{keyword}'")
                    continue
                
                logger.debug(f"   🔍 {len(search_results)} résultats Serper pour '{keyword}'")
                
                for result in search_results:
                    if pages_visited >= self.max_pages_to_visit:
                        break
                    
                    url = result.get('link') or result.get('url')
                    if not url or url in visited_urls:
                        continue
                    
                    # Filtrer les URLs non pertinentes
                    if not self._is_valid_url(url):
                        continue
                    
                    visited_urls.add(url)
                    pages_visited += 1
                    
                    logger.info(f"   📄 [{pages_visited}/{self.max_pages_to_visit}] Visit: {url[:70]}...")
                    
                    # Analyser la page (télécharger + parser)
                    page_data = self._analyze_page(url)
                    if page_data and self._is_relevant_content(page_data):
                        
                        # Extraire les informations structurées
                        offre = self._extract_offre_data(url, page_data, keyword)
                        if offre:
                            all_offres.append(offre)
                            
                            # Télécharger les documents associés si activé
                            if self.download_documents:
                                docs = self._download_documents(url, page_data)
                                if docs:
                                    offre['attached_documents'] = docs
                                    logger.info(f"   📎 {len(docs)} document(s) téléchargé(s)")
                    
                    # Delay entre les pages pour respecter les serveurs
                    time.sleep(self.delay_between_requests)
                
                # Delay entre les mots-clés
                time.sleep(self.delay_between_keywords)
                
            except Exception as e:
                logger.error(f"❌ Erreur recherche '{keyword}': {e}")
                continue
        
        logger.info(f"📦 WebSearch Serper terminé: {len(all_offres)} offres trouvées")
        logger.info(f"   📊 Stats: {pages_visited} pages visitées, {api_calls} appels API Serper")
        
        if all_offres:
            self.save_to_db(all_offres)
            logger.info(f"   💾 {len(all_offres)} offres sauvegardées dans MongoDB")
        
        return all_offres
    
    def _search_serper(self, query: str, max_results: int = 15) -> List[Dict]:
        """Effectue une recherche via l'API Serper.dev"""
        headers = {
            'X-API-KEY': self.api_key,
            'Content-Type': 'application/json',
            'User-Agent': self._get_random_user_agent()
        }
        
        payload = {
            'q': query,
            'num': max_results,
            'gl': 'ci',      # Géolocalisation: Côte d'Ivoire
            'hl': 'fr',      # Langue: Français
            'tbs': 'qdr:y'   # Résultats de l'année dernière
        }
        
        try:
            response = requests.post(
                self.api_url,
                headers=headers,
                json=payload,
                timeout=30
            )
            response.raise_for_status()
            
            data = response.json()
            
            # Extraire les résultats organiques
            results = []
            for item in data.get('organic', []):
                results.append({
                    'title': item.get('title', ''),
                    'link': item.get('link', ''),
                    'snippet': item.get('snippet', ''),
                    'source': item.get('source', ''),
                    'date': item.get('date', ''),
                    'position': item.get('position', 0)
                })
            
            return results
            
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 401:
                logger.error("❌ Clé API Serper invalide")
            elif e.response.status_code == 429:
                logger.error("❌ Limite de requêtes Serper dépassée")
            else:
                logger.error(f"❌ Erreur HTTP Serper: {e}")
            return []
        except Exception as e:
            logger.error(f"❌ Erreur requête Serper: {e}")
            return []
    
    def _is_valid_url(self, url: str) -> bool:
        """Filtre les URLs non pertinentes ou dangereuses"""
        try:
            parsed = urlparse(url)
            
            # Accepter uniquement http/https
            if parsed.scheme not in ['http', 'https']:
                return False
            
            domain = parsed.netloc.lower()
            
            # Exclure les réseaux sociaux (contenu non structuré)
            excluded_domains = [
                'linkedin.com', 'facebook.com', 'twitter.com', 'x.com',
                'instagram.com', 'youtube.com', 'tiktok.com', 'pinterest.com'
            ]
            if any(excl in domain for excl in excluded_domains):
                return False
            
            # Exclure les URLs avec patterns d'emploi
            excluded_patterns = ['/job', '/career', '/emploi', '/recrutement', '/hiring', '/postuler']
            if any(pat in url.lower() for pat in excluded_patterns):
                return False
            
            # Exclure les URLs trop courtes ou suspectes
            if len(url) < 20 or 'javascript:' in url.lower():
                return False
            
            return True
            
        except Exception:
            return False
    
    def _analyze_page(self, url: str) -> Optional[Dict]:
        """Télécharge et analyse le contenu d'une page"""
        try:
            headers = {
                'User-Agent': self._get_random_user_agent(),
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'fr-FR,fr;q=0.9,en;q=0.8',
                'Connection': 'keep-alive',
            }
            
            response = requests.get(url, headers=headers, timeout=25, allow_redirects=True)
            response.raise_for_status()
            
            # Vérifier le type de contenu
            content_type = response.headers.get('Content-Type', '').lower()
            
            # Si c'est un document direct (PDF, Word, etc.)
            if any(doc_type in content_type for doc_type in ['pdf', 'msword', 'excel', 'csv', 'vnd.openxmlformats']):
                return {
                    'is_document': True,
                    'content_type': content_type,
                    'content': response.content,
                    'text': self._parse_binary_content(response.content, content_type),
                    'links': [],
                    'title': Path(url).name or 'Document',
                    'url': url
                }
            
            # Parser le HTML
            soup = self._parse_html(response.text)
            if not soup:
                return None
            
            # Extraire le texte principal
            text_content = self._extract_main_text(soup)
            if len(text_content.strip()) < 50:  # Page trop vide
                return None
            
            # Extraire les liens vers des documents
            document_links = self._find_document_links(soup, url)
            
            # Extraire les métadonnées
            title = self._get_meta_tag(soup, 'title') or soup.find('title')
            title_text = title.text.strip() if title else Path(url).name
            
            return {
                'is_document': False,
                'content_type': content_type,
                'text': text_content[:15000],  # Limiter à 15k caractères
                'links': document_links,
                'title': title_text[:200],
                'meta_description': self._get_meta_tag(soup, 'description')[:300],
                'meta_keywords': self._get_meta_tag(soup, 'keywords'),
                'url': url,
                'response_status': response.status_code
            }
            
        except requests.exceptions.Timeout:
            logger.debug(f"   ⏰ Timeout chargement page: {url[:60]}")
            return None
        except requests.exceptions.RequestException as e:
            logger.debug(f"   ❌ Erreur HTTP: {e}")
            return None
        except Exception as e:
            logger.debug(f"   ❌ Erreur analyse page: {e}")
            return None
    
    def _extract_main_text(self, soup) -> str:
        """Extrait le texte principal d'une page HTML"""
        # Supprimer les éléments non pertinents
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()
        
        # Chercher le contenu principal par ordre de priorité
        main_content = (
            soup.find('main') or 
            soup.find('article') or 
            soup.find('div', class_='content') or
            soup.find('div', id='content') or
            soup.find('div', class_='article') or
            soup.body
        )
        
        if main_content:
            text = main_content.get_text(separator='\n', strip=True)
            # Nettoyer: espaces multiples, lignes vides
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r'[ \t]+', ' ', text)
            return text.strip()[:12000]  # Limiter à 12k caractères
        
        return soup.get_text(separator=' ', strip=True)[:5000]
    
    def _find_document_links(self, soup, base_url: str) -> List[Dict]:
        """Trouve les liens vers des documents téléchargeables"""
        documents = []
        seen_urls = set()
        
        for link in soup.find_all('a', href=True):
            href = link['href'].strip()
            if not href or href.startswith('#') or href.startswith('javascript:'):
                continue
            
            text = link.get_text(strip=True).lower()
            
            # Vérifier si c'est un lien vers un document par extension
            if any(href.lower().split('?')[0].endswith(ext) for ext in self.document_extensions):
                full_url = urljoin(base_url, href)
                if full_url not in seen_urls:
                    seen_urls.add(full_url)
                    documents.append({
                        'url': full_url,
                        'text': text or Path(href).name,
                        'extension': Path(href).suffix.lower().split('?')[0]
                    })
            
            # Ou si le texte du lien indique un document
            elif any(kw in text for kw in ['pdf', 'télécharger', 'download', 'doc', 'excel', 'dossier']):
                full_url = urljoin(base_url, href)
                ext = self._detect_extension_from_url(full_url)
                if ext and full_url not in seen_urls:
                    seen_urls.add(full_url)
                    documents.append({
                        'url': full_url,
                        'text': text,
                        'extension': ext
                    })
            
            # Limiter à 10 documents max par page
            if len(documents) >= 10:
                break
        
        return documents
    
    def _detect_extension_from_url(self, url: str) -> Optional[str]:
        """Détecte l'extension de fichier depuis une URL"""
        parsed = urlparse(url)
        path = parsed.path.lower().split('?')[0]  # Supprimer les paramètres
        
        for ext in self.document_extensions:
            if path.endswith(ext):
                return ext
        return None
    
    def _is_relevant_content(self, page_data: Dict) -> bool:
        """Détermine si le contenu concerne un appel d'offre"""
        text = (
            page_data.get('title', '') + ' ' + 
            page_data.get('text', '') + ' ' + 
            page_data.get('meta_description', '') + ' ' +
            page_data.get('meta_keywords', '')
        ).lower()
        
        if len(text.strip()) < 100:
            return False
        
        # Doit contenir au moins un mot-clé pertinent
        has_relevant = any(kw in text for kw in self.relevant_keywords)
        if not has_relevant:
            return False
        
        # Ne doit pas contenir de mots-clés d'exclusion
        has_excluded = any(kw in text for kw in self.exclude_keywords)
        if has_excluded:
            return False
        
        return True
    
    def _extract_offre_data(self, url: str, page_data: Dict, keyword: str) -> Optional[Dict]:
        """Extrait les données structurées d'un appel d'offre"""
        text = page_data.get('text', '')
        title = page_data.get('title', '')
        
        # Titre
        offre_title = title or self._extract_title_from_text(text) or keyword
        
        # Description (snippet pertinent)
        description = self._extract_snippet(text, max_length=500)
        
        # Date de publication (approximative)
        date_pub = self._extract_date_from_content(text)
        
        # Localisation
        location = self._extract_location(text)
        
        # Type d'offre
        employment_type = self._classify_offer_type(text)
        
        # Entreprise/organisme
        company = self._extract_organization(text)
        
        return {
            'title': offre_title[:200],
            'url': url,
            'source_url': url,
            'description': description,
            'category': 'Web Search',
            'subcategory': keyword,
            'location': location or 'Côte d\'Ivoire',
            'employment_type': employment_type,
            'date_publication': date_pub,
            'company_name': company,
            'tags': ['serper', 'web-search', keyword, 'auto-detect'],
            'source_type': 'serper_api',
            'search_keyword': keyword,
            'content_preview': text[:1000],
            'meta_description': page_data.get('meta_description', ''),
            'scraped_at': datetime.utcnow().isoformat()
        }
    
    def _extract_title_from_text(self, text: str) -> Optional[str]:
        """Extrait un titre potentiel depuis le texte"""
        lines = [l.strip() for l in text.split('\n') if 20 < len(l.strip()) < 150]
        
        for line in lines[:15]:  # Chercher dans les premières lignes pertinentes
            line_lower = line.lower()
            if any(kw in line_lower for kw in self.relevant_keywords):
                # Nettoyer le titre
                title = re.sub(r'\s*[-–|•]\s*.*$', '', line)  # Supprimer suffixes
                return title.strip()[:150]
        return None
    
    def _extract_snippet(self, text: str, max_length: int = 500) -> str:
        """Extrait un snippet pertinent du texte"""
        # Chercher les phrases contenant des mots-clés
        sentences = re.split(r'[.!?]+', text)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if 30 < len(sentence) < max_length:
                if any(kw in sentence.lower() for kw in self.relevant_keywords):
                    return sentence + ('...' if len(sentence) >= max_length - 3 else '')
        
        # Fallback: premières lignes non vides
        lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 20]
        return ' '.join(lines[:3])[:max_length]
    
    def _extract_date_from_content(self, text: str) -> datetime:
        """Extrait une date depuis le contenu"""
        patterns = [
            # Formats français
            (r'(\d{1,2}/\d{1,2}/\d{4})', lambda x: self._normalize_date(x)),
            (r'(\d{1,2}\s+[a-zA-Z]+\s+\d{4})', lambda x: self._normalize_date(x)),
            (r'publié\s+(?:le\s+)?([\d\s/a-zA-Z]+)', lambda x: self._normalize_date(x)),
            (r'date\s*[:\s]+([\d\s/a-zA-Z]+)', lambda x: self._normalize_date(x)),
            # Formats ISO/US
            (r'(\d{4}-\d{1,2}-\d{1,2})', lambda x: self._normalize_date(x)),
            # Relatif
            (r'il y a\s+(\d+)\s+jours?', lambda x: datetime.utcnow() - timedelta(days=int(x))),
            (r'il y a\s+(\d+)\s+semaines?', lambda x: datetime.utcnow() - timedelta(weeks=int(x))),
        ]
        
        for pattern, converter in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return converter(match.group(1))
                except:
                    continue
        
        return datetime.utcnow()
    
    def _extract_location(self, text: str) -> Optional[str]:
        """Extrait la localisation depuis le texte"""
        locations = {
            'abidjan': 'Abidjan',
            'bouaké': 'Bouaké', 'bouake': 'Bouaké',
            'yamoussoukro': 'Yamoussoukro',
            'san-pédro': 'San-Pédro', 'san pedro': 'San-Pédro',
            'korhogo': 'Korhogo',
            'daloa': 'Daloa',
            'man': 'Man',
            'gagnoa': 'Gagnoa',
            'divo': 'Divo',
            'bondoukou': 'Bondoukou',
            'côte d\'ivoire': 'Côte d\'Ivoire',
            'cote d\'ivoire': 'Côte d\'Ivoire',
            'ivory coast': 'Côte d\'Ivoire',
            'ci': 'Côte d\'Ivoire'
        }
        
        text_lower = text.lower()
        for key, value in locations.items():
            # Vérifier que c'est un mot complet (pas "ci" dans "microsoft")
            if re.search(r'\b' + re.escape(key) + r'\b', text_lower):
                return value
        return None
    
    def _classify_offer_type(self, text: str) -> str:
        """Classifie le type d'appel d'offre"""
        text_lower = text.lower()
        
        classifications = [
            (['informatique', 'réseau', 'cyber', 'microsoft', 'application', 'système', 'logiciel'], 'Appel d\'offre - Informatique'),
            (['batiment', 'construction', 'travaux', 'infrastructure', 'génie civil', 'btp'], 'Appel d\'offre - BTP'),
            (['matériel', 'équipement', 'fourniture', 'acquisition', 'achat'], 'Appel d\'offre - Fournitures'),
            (['consultance', 'consulting', 'étude', 'expertise', 'conseil', 'audit'], 'Appel d\'offre - Consultance'),
            (['formation', 'training', 'capacitation', 'renforcement'], 'Appel d\'offre - Formation'),
            (['service', 'prestation', 'maintenance', 'support', 'assistance'], 'Appel d\'offre - Services'),
        ]
        
        for keywords, category in classifications:
            if any(kw in text_lower for kw in keywords):
                return category
        
        return 'Appel d\'offre'
    
    def _extract_organization(self, text: str) -> Optional[str]:
        """Tente d'extraire le nom de l'organisme"""
        # Patterns pour organismes en Côte d'Ivoire
        patterns = [
            r'(?:ministère|direction|agence|office|société|établissement)\s+[^.\n]{10,60}',
            r'(?:république\s+de\s+)?c[oô]te\s+d\'?ivoire[^.\n]{0,30}',
            r'[A-Z]{2,}(?:\s+[A-Z][a-z]+){2,}',  # Acronymes comme ARMP, PRICI
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                org = match.group(0).strip()
                # Validation basique
                if 10 < len(org) < 100 and org.lower() not in ['côte d\'ivoire', 'république de côte d\'ivoire']:
                    return org
        return None
    
    def _download_documents(self, page_url: str, page_data: Dict) -> List[Dict]:
        """Télécharge et parse les documents trouvés"""
        if not DOCUMENT_PARSING_AVAILABLE:
            logger.debug("   ⚠️  Parsing de documents désactivé (libs manquantes)")
            return []
        
        downloaded = []
        document_links = page_data.get('links', [])
        
        for doc in document_links:
            try:
                url = doc['url']
                ext = doc['extension']
                
                logger.debug(f"   📥 Téléchargement: {url[:60]}...")
                
                # Télécharger le fichier
                response = requests.get(url, timeout=40, stream=True, headers={
                    'User-Agent': self._get_random_user_agent()
                })
                response.raise_for_status()
                
                # Nommer le fichier de façon unique
                filename = self._sanitize_filename(doc['text'] or Path(url).name)
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                filepath = self.download_folder / f"{timestamp}_{filename}"
                
                # Sauvegarder le fichier
                with open(filepath, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                
                # Parser le contenu si possible
                content_text = self._parse_document(filepath, ext)
                
                downloaded.append({
                    'filename': filename,
                    'filepath': str(filepath),
                    'url': url,
                    'extension': ext,
                    'size_bytes': os.path.getsize(filepath),
                    'parsed_content_preview': content_text[:500] if content_text else None,
                    'downloaded_at': datetime.utcnow().isoformat()
                })
                
                logger.debug(f"   ✅ Document téléchargé: {filename} ({os.path.getsize(filepath)} bytes)")
                
            except requests.exceptions.RequestException as e:
                logger.debug(f"   ❌ Erreur téléchargement {url[:50]}: {e}")
                continue
            except Exception as e:
                logger.debug(f"   ❌ Erreur traitement document: {e}")
                continue
        
        return downloaded
    
    def _parse_binary_content(self, content: bytes, content_type: str) -> str:
        """Parse le contenu binaire d'un document"""
        if not DOCUMENT_PARSING_AVAILABLE:
            return ""
        
        try:
            if 'pdf' in content_type and 'PyPDF2' in globals():
                from io import BytesIO
                reader = PyPDF2.PdfReader(BytesIO(content))
                text = ''
                for page in reader.pages[:10]:  # Limiter aux 10 premières pages
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                return text[:10000]
            
            elif 'msword' in content_type or 'openxmlformats-officedocument.wordprocessingml' in content_type:
                from io import BytesIO
                doc = WordDocument(BytesIO(content))
                return '\n'.join([p.text for p in doc.paragraphs])[:10000]
            
        except Exception as e:
            logger.debug(f"   Erreur parsing binaire: {e}")
        
        return ""
    
    def _parse_document(self, filepath: Path, ext: str) -> Optional[str]:
        """Extrait le texte d'un document local"""
        if not DOCUMENT_PARSING_AVAILABLE:
            return None
        
        try:
            if ext == '.pdf' and 'PyPDF2' in globals():
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''
                    for page in reader.pages[:10]:  # Limiter aux 10 premières pages
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                    return text[:10000]
            
            elif ext in ['.doc', '.docx'] and 'WordDocument' in globals():
                doc = WordDocument(str(filepath))
                return '\n'.join([p.text for p in doc.paragraphs])[:10000]
            
            elif ext in ['.xls', '.xlsx'] and 'openpyxl' in globals():
                wb = openpyxl.load_workbook(filepath, data_only=True)
                texts = []
                for sheet in wb.worksheets[:3]:  # 3 premières feuilles
                    for row in sheet.iter_rows(values_only=True):
                        row_text = ' | '.join(str(c) for c in row if c is not None)
                        if row_text.strip():
                            texts.append(row_text)
                return '\n'.join(texts[:100])[:10000]
            
            elif ext == '.csv' and 'pandas' in globals():
                df = pd.read_csv(filepath, nrows=50, on_bad_lines='skip')
                return df.to_string()[:10000]
            
        except Exception as e:
            logger.debug(f"   Erreur parsing {filepath.name}: {e}")
        
        return None
    
    def _sanitize_filename(self, filename: str) -> str:
        """Nettoie un nom de fichier pour le stockage sécurisé"""
        # Supprimer les caractères invalides pour Windows/Linux
        filename = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', filename)
        # Supprimer les espaces multiples
        filename = re.sub(r'\s+', '_', filename)
        # Limiter la longueur
        name, ext = os.path.splitext(filename)
        return f"{name[:80]}{ext}" or 'document'
    
    def _get_meta_tag(self, soup, name: str) -> str:
        """Extrait une balise meta par nom ou property"""
        meta = (soup.find('meta', attrs={'name': name}) or 
                soup.find('meta', attrs={'property': name}) or
                soup.find('meta', attrs={'itemprop': name}))
        return meta.get('content', '').strip() if meta else ''
    
    def _get_random_user_agent(self) -> str:
        """Retourne un User-Agent réaliste et varié"""
        return random.choice([
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36',
        ])